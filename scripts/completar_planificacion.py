#!/usr/bin/env python3
"""
completar_planificacion.py — Completa documentos de planificación curricular
Autor: proscar.cl · https://github.com/oscraft2/skill-codex-claude

Uso:
    python scripts/completar_planificacion.py --input formato.docx --asignatura "Matemática" --nivel "1° Medio"
    python scripts/completar_planificacion.py --tipo anual --asignatura "Matemática" --nivel "2° Medio"
    python scripts/completar_planificacion.py --tipo unidad --asignatura "Matemática" --nivel "3° Medio" --unidad 1

Requiere:
    pip install python-docx
"""

import json, argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
    print("⚠️  Instala python-docx: pip install python-docx")

BASE_DIR = Path(__file__).parent.parent
OUT_DIR  = BASE_DIR / "output"
OUT_DIR.mkdir(exist_ok=True)

# ── Colores institucionales ───────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0, 70, 127)
AZUL_CLARO   = RGBColor(173, 216, 230)
VERDE        = RGBColor(0, 112, 80)
GRIS         = RGBColor(240, 240, 240)

# ── OA por nivel y asignatura (fuente: curriculumnacional.cl) ─────────────────
OA_DATABASE = {
    "Matemática": {
        "1° Medio": {
            "ejes": ["Números", "Álgebra y funciones", "Geometría", "Probabilidad y estadística"],
            "unidades": [
                {
                    "nombre": "Números racionales y potencias",
                    "eje": "Números",
                    "oa": [
                        {"codigo":"MA1M OA 01","descripcion":"Calcular operaciones con números racionales en forma simbólica."},
                        {"codigo":"MA1M OA 02","descripcion":"Mostrar que comprenden las potencias de base racional y exponente entero."},
                    ],
                    "horas": 20
                },
                {
                    "nombre": "Álgebra: productos notables y sistemas",
                    "eje": "Álgebra y funciones",
                    "oa": [
                        {"codigo":"MA1M OA 03","descripcion":"Desarrollar los productos notables de manera concreta, pictórica y simbólica."},
                        {"codigo":"MA1M OA 04","descripcion":"Resolver sistemas de ecuaciones lineales (2x2) relacionados con problemas de la vida diaria."},
                        {"codigo":"MA1M OA 05","descripcion":"Graficar relaciones lineales en dos variables de la forma f(x,y)=ax+by."},
                    ],
                    "horas": 30
                },
                {
                    "nombre": "Geometría: homotecia y semejanza",
                    "eje": "Geometría",
                    "oa": [
                        {"codigo":"MA1M OA 06","descripcion":"Desarrollar la fórmula del área y perímetro de sectores y segmentos circulares."},
                        {"codigo":"MA1M OA 07","descripcion":"Desarrollar las fórmulas para el área de superficie y volumen del cono."},
                        {"codigo":"MA1M OA 08","descripcion":"Mostrar que comprenden el concepto de homotecia."},
                        {"codigo":"MA1M OA 09","descripcion":"Desarrollar el teorema de Tales mediante las propiedades de la homotecia."},
                        {"codigo":"MA1M OA 10","descripcion":"Aplicar propiedades de semejanza y proporcionalidad a modelos a escala."},
                        {"codigo":"MA1M OA 11","descripcion":"Representar el concepto de homotecia de forma vectorial."},
                    ],
                    "horas": 40
                },
                {
                    "nombre": "Probabilidad y estadística",
                    "eje": "Probabilidad y estadística",
                    "oa": [
                        {"codigo":"MA1M OA 12","descripcion":"Registrar distribuciones en tabla de doble entrada y nube de puntos."},
                        {"codigo":"MA1M OA 13","descripcion":"Comparar poblaciones mediante gráficos xy para dos atributos de muestras."},
                        {"codigo":"MA1M OA 14","descripcion":"Desarrollar las reglas de las probabilidades: aditiva, multiplicativa y combinación."},
                        {"codigo":"MA1M OA 15","descripcion":"Mostrar que comprenden el concepto de azar mediante experimentación y análisis."},
                    ],
                    "horas": 30
                }
            ]
        },
        "2° Medio": {
            "ejes": ["Números", "Álgebra y funciones", "Geometría", "Probabilidad y estadística"],
            "unidades": [
                {
                    "nombre": "Números reales y raíces",
                    "eje": "Números",
                    "oa": [
                        {"codigo":"MA2M OA 01","descripcion":"Realizar cálculos con números reales: descomposición de raíces y propiedades."},
                        {"codigo":"MA2M OA 02","descripcion":"Mostrar que comprenden los números reales y su representación en la recta numérica."},
                    ],
                    "horas": 20
                },
                {
                    "nombre": "Función cuadrática e inversa",
                    "eje": "Álgebra y funciones",
                    "oa": [
                        {"codigo":"MA2M OA 03","descripcion":"Mostrar que comprenden la función cuadrática f(x)=ax²+bx+c."},
                        {"codigo":"MA2M OA 04","descripcion":"Resolver ecuaciones cuadráticas de diversas formas."},
                        {"codigo":"MA2M OA 05","descripcion":"Mostrar que comprenden la inversa de una función."},
                        {"codigo":"MA2M OA 06","descripcion":"Explicar el cambio porcentual constante en intervalos de tiempo."},
                    ],
                    "horas": 35
                },
                {
                    "nombre": "Geometría: esfera y trigonometría",
                    "eje": "Geometría",
                    "oa": [
                        {"codigo":"MA2M OA 07","descripcion":"Desarrollar las fórmulas del área de superficie y volumen de la esfera."},
                        {"codigo":"MA2M OA 08","descripcion":"Mostrar que comprenden las razones trigonométricas sen, cos y tan en triángulos rectángulos."},
                        {"codigo":"MA2M OA 09","descripcion":"Aplicar las razones trigonométricas en composición y descomposición de vectores."},
                    ],
                    "horas": 30
                },
                {
                    "nombre": "Probabilidad y estadística",
                    "eje": "Probabilidad y estadística",
                    "oa": [
                        {"codigo":"MA2M OA 10","descripcion":"Mostrar que comprenden las variables aleatorias finitas."},
                        {"codigo":"MA2M OA 11","descripcion":"Utilizar permutaciones y combinatoria para calcular probabilidades."},
                        {"codigo":"MA2M OA 12","descripcion":"Mostrar que comprenden el rol de la probabilidad en la sociedad."},
                    ],
                    "horas": 35
                }
            ]
        },
        "3° Medio": {
            "ejes": ["Números y Álgebra", "Álgebra y funciones", "Geometría", "Probabilidad y estadística"],
            "unidades": [
                {
                    "nombre": "Números complejos",
                    "eje": "Números y Álgebra",
                    "oa": [
                        {"codigo":"FG-MATE-3M-OAC-01","descripcion":"Resolver problemas de adición, sustracción, multiplicación y división de números complejos."},
                    ],
                    "horas": 30
                },
                {
                    "nombre": "Estadística y probabilidad condicional",
                    "eje": "Probabilidad y estadística",
                    "oa": [
                        {"codigo":"FG-MATE-3M-OAC-02","descripcion":"Tomar decisiones en situaciones de incerteza con medidas de dispersión y probabilidades condicionales."},
                    ],
                    "horas": 30
                },
                {
                    "nombre": "Funciones exponencial y logarítmica",
                    "eje": "Álgebra y funciones",
                    "oa": [
                        {"codigo":"FG-MATE-3M-OAC-03","descripcion":"Aplicar modelos matemáticos con funciones exponencial y logarítmica."},
                    ],
                    "horas": 35
                },
                {
                    "nombre": "Geometría de la circunferencia",
                    "eje": "Geometría",
                    "oa": [
                        {"codigo":"FG-MATE-3M-OAC-04","descripcion":"Resolver problemas de geometría euclidiana con relaciones métricas en la circunferencia."},
                    ],
                    "horas": 25
                }
            ]
        },
        "4° Medio": {
            "ejes": ["Números y economía", "Álgebra y funciones", "Geometría", "Probabilidad y estadística"],
            "unidades": [
                {
                    "nombre": "Matemática financiera",
                    "eje": "Números y economía",
                    "oa": [
                        {"codigo":"FG-MATE-4M-OAC-01","descripcion":"Fundamentar decisiones en el ámbito financiero usando modelos con porcentajes, tasas e índices económicos."},
                    ],
                    "horas": 30
                },
                {
                    "nombre": "Estadística y modelos probabilísticos",
                    "eje": "Probabilidad y estadística",
                    "oa": [
                        {"codigo":"FG-MATE-4M-OAC-02","descripcion":"Fundamentar decisiones con análisis crítico de datos estadísticos y modelos binomial y normal."},
                    ],
                    "horas": 35
                },
                {
                    "nombre": "Funciones trigonométricas y potencias",
                    "eje": "Álgebra y funciones",
                    "oa": [
                        {"codigo":"FG-MATE-4M-OAC-03","descripcion":"Construir modelos con funciones potencias de exponente entero y trigonométricas sen(x) y cos(x)."},
                    ],
                    "horas": 35
                },
                {
                    "nombre": "Geometría analítica",
                    "eje": "Geometría analítica",
                    "oa": [
                        {"codigo":"FG-MATE-4M-OAC-04","descripcion":"Resolver problemas de rectas y circunferencias en el plano mediante representación analítica."},
                    ],
                    "horas": 20
                }
            ]
        }
    }
}

HABILIDADES_MATE = {
    "1° Medio": ["Resolver problemas", "Argumentar y comunicar", "Modelar", "Representar"],
    "2° Medio": ["Resolver problemas", "Argumentar y comunicar", "Modelar", "Representar"],
    "3° Medio": ["Construir y evaluar estrategias colaborativas", "Variar parámetros en modelos", "Tomar decisiones con evidencia estadística", "Argumentar con lenguaje simbólico", "Construir modelos predictivos"],
    "4° Medio": ["Construir y evaluar estrategias colaborativas", "Variar parámetros en modelos", "Tomar decisiones con evidencia estadística", "Argumentar con lenguaje simbólico", "Construir modelos predictivos"],
}

INDICADORES_BASE = {
    "Resolver problemas": "Aplican estrategias de resolución (descomponer, buscar patrones, estimar) para abordar situaciones nuevas.",
    "Argumentar y comunicar": "Justifican sus procedimientos usando lenguaje matemático, esquemas y gráficos.",
    "Modelar": "Seleccionan y ajustan modelos matemáticos para representar fenómenos del entorno.",
    "Representar": "Transitan entre distintas representaciones (concreta, pictórica y simbólica) de un mismo concepto.",
}


def set_cell_color(cell, color: RGBColor):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def set_cell_bold_white(cell, texto: str, size=11, center=False):
    para = cell.paragraphs[0]
    para.clear()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(texto)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(255, 255, 255)


def generar_planificacion_anual(asignatura: str, nivel: str, output: str):
    """Genera una planificación anual en formato DOCX."""
    if not DOCX_OK:
        return

    datos = OA_DATABASE.get(asignatura, {}).get(nivel)
    if not datos:
        print(f"⚠️  No hay datos para {asignatura} {nivel}. Agrega OA al script.")
        return

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── Título ────────────────────────────────────────────────────────────────
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run(f"PLANIFICACIÓN ANUAL\n{asignatura.upper()} — {nivel}")
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = AZUL_OSCURO

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitulo.add_run(f"Año escolar {datetime.now().year} · Bases Curriculares Mineduc")
    run2.font.size  = Pt(10)
    run2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # ── Tabla de identificación ───────────────────────────────────────────────
    t_id = doc.add_table(rows=2, cols=4)
    t_id.style = "Table Grid"
    encabezados_id = ["Establecimiento:", "Asignatura:", "Nivel:", "Docente:"]
    valores_id     = ["", asignatura, nivel, ""]
    for i, (enc, val) in enumerate(zip(encabezados_id, valores_id)):
        set_cell_color(t_id.rows[0].cells[i], AZUL_OSCURO)
        set_cell_bold_white(t_id.rows[0].cells[i], enc, size=10, center=True)
        t_id.rows[1].cells[i].text = val

    doc.add_paragraph()

    # ── Tabla principal por unidad ────────────────────────────────────────────
    encabezados = [
        "Unidad", "Eje temático", "Objetivos de Aprendizaje",
        "Habilidades", "Horas\npedag.", "Evaluación"
    ]
    tabla = doc.add_table(rows=1, cols=6)
    tabla.style = "Table Grid"

    # Fila de encabezado
    hdr = tabla.rows[0]
    for i, enc in enumerate(encabezados):
        set_cell_color(hdr.cells[i], AZUL_OSCURO)
        set_cell_bold_white(hdr.cells[i], enc, size=10, center=True)

    habilidades = HABILIDADES_MATE.get(nivel, ["Resolver problemas", "Modelar", "Representar"])

    for idx, unidad in enumerate(datos["unidades"], 1):
        row = tabla.add_row()

        # Unidad
        row.cells[0].text = f"Unidad {idx}\n{unidad['nombre']}"
        row.cells[0].paragraphs[0].runs[0].bold = True

        # Eje
        row.cells[1].text = unidad["eje"]

        # OA
        oa_texto = "\n".join([f"• {oa['codigo']}: {oa['descripcion']}" for oa in unidad["oa"]])
        row.cells[2].text = oa_texto

        # Habilidades
        row.cells[3].text = "\n".join([f"• {h}" for h in habilidades[:3]])

        # Horas
        row.cells[4].text = str(unidad["horas"])
        row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Evaluación
        if idx == 1:
            eval_tipo = "Diagnóstica\nFormativa"
        elif idx == len(datos["unidades"]):
            eval_tipo = "Formativa\nSumativa"
        else:
            eval_tipo = "Formativa\nSumativa"
        row.cells[5].text = eval_tipo

    doc.add_paragraph()

    # ── Nota al pie ───────────────────────────────────────────────────────────
    nota = doc.add_paragraph()
    run_nota = nota.add_run(
        "OA verificables en curriculumnacional.cl · "
        "Generado con planificacion-curricular-chile skill · proscar.cl"
    )
    run_nota.font.size      = Pt(8)
    run_nota.font.color.rgb = RGBColor(150, 150, 150)
    nota.alignment          = WD_ALIGN_PARAGRAPH.CENTER

    # Guardar
    out_path = Path(output)
    doc.save(out_path)
    print(f"✅ Planificación anual: {out_path}")


def completar_docx(input_path: str, asignatura: str, nivel: str):
    """
    Analiza un DOCX existente e intenta completar campos vacíos
    basándose en los OA del nivel/asignatura indicados.
    """
    if not DOCX_OK:
        return

    doc = Document(input_path)
    datos = OA_DATABASE.get(asignatura, {}).get(nivel)

    if not datos:
        print(f"⚠️  Sin datos para {asignatura} {nivel}. Generando planificación nueva.")
        out = str(OUT_DIR / f"planificacion_{nivel.replace('°','').replace(' ','_')}_{asignatura[:3]}.docx")
        generar_planificacion_anual(asignatura, nivel, out)
        return

    # Buscar celdas vacías en tablas y completarlas con OA
    oa_todos = []
    for unidad in datos["unidades"]:
        for oa in unidad["oa"]:
            oa_todos.append(f"{oa['codigo']}: {oa['descripcion']}")

    completados = 0
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                texto = celda.text.strip().lower()
                # Detectar campos típicos vacíos
                if texto in ["", "oa", "objetivo", "objetivos", "obj. de aprendizaje"]:
                    if oa_todos:
                        celda.text = oa_todos.pop(0)
                        completados += 1

    out_path = Path(input_path).stem + f"_completado_{nivel.replace('°','').replace(' ','_')}.docx"
    out_path = OUT_DIR / out_path
    doc.save(out_path)
    print(f"✅ Documento completado ({completados} campos): {out_path}")
    print(f"   Revisa y ajusta los campos completados según tu formato institucional.")


def main():
    ap = argparse.ArgumentParser(description="Completa planificaciones curriculares chilenas")
    ap.add_argument("--input",       help="DOCX existente a completar")
    ap.add_argument("--tipo",        default="anual", choices=["anual","unidad","clase"])
    ap.add_argument("--asignatura",  default="Matemática")
    ap.add_argument("--nivel",       default="1° Medio")
    ap.add_argument("--unidad",      type=int, default=1)
    ap.add_argument("--output",      default="")
    args = ap.parse_args()

    if not DOCX_OK:
        print("Instala python-docx: pip install python-docx")
        return

    if args.input:
        completar_docx(args.input, args.asignatura, args.nivel)
    else:
        nombre = args.output or str(
            OUT_DIR / f"planificacion_{args.tipo}_{args.nivel.replace('°','').replace(' ','_')}_{args.asignatura[:3]}.docx"
        )
        generar_planificacion_anual(args.asignatura, args.nivel, nombre)


if __name__ == "__main__":
    main()
