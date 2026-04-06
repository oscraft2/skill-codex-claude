#!/usr/bin/env python3
"""
generador_filas.py — Genera filas B, C, D de una prueba a partir de la fila A
Autor: proscar.cl · https://github.com/oscraft2/skill-codex-claude

Uso:
    python scripts/generador_filas.py --input prueba_A.docx --filas B C D
    python scripts/generador_filas.py --input prueba_A.docx --filas B --semilla 99
    python scripts/generador_filas.py --input prueba_A.docx --filas B C --no-mezclar-alternativas

Requiere: pip install python-docx
"""

import argparse, copy, random, re, json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy as _copy
except ImportError:
    print("Instala python-docx: pip install python-docx")
    raise SystemExit(1)

OUT_DIR = Path(__file__).parent.parent / "output"
OUT_DIR.mkdir(exist_ok=True)


# ── Colores ───────────────────────────────────────────────────────────────────
AZUL   = RGBColor(0, 70, 127)
BLANCO = RGBColor(255, 255, 255)

# ── Patrones de detección ─────────────────────────────────────────────────────
RE_PREGUNTA = re.compile(r'^(\d{1,2})[.)]\s+(.+)', re.DOTALL)
RE_ALT      = re.compile(r'^([a-eA-E])[.)]\s+(.+)', re.DOTALL)


# ─────────────────────────────────────────────────────────────────────────────
# LECTOR: extrae preguntas del .docx
# ─────────────────────────────────────────────────────────────────────────────

class Pregunta:
    def __init__(self, numero, enunciado, alternativas, respuesta_correcta=None):
        self.numero          = numero           # número original (1, 2, 3...)
        self.enunciado       = enunciado        # texto del enunciado
        self.alternativas    = alternativas     # lista de (letra, texto)
        self.respuesta_correcta = respuesta_correcta  # letra correcta (si se conoce)

    def __repr__(self):
        return f"P{self.numero}: {self.enunciado[:40]}... [{len(self.alternativas)} alts]"


def leer_preguntas_docx(path: str) -> tuple:
    """
    Lee un .docx y extrae:
    - lista de Pregunta
    - párrafos de encabezado (antes de la primera pregunta)
    - párrafos de pie (después de la última alternativa)
    """
    doc = Document(path)
    preguntas = []
    encabezado_parrs = []  # párrafos antes de la primera pregunta
    pie_parrs = []         # párrafos después de todo

    # Extraer todo el texto como lista de párrafos con sus objetos
    todos = list(doc.paragraphs)

    # También extraer párrafos de tablas (para formato ISPM de 2 columnas)
    parrs_con_tabla = []
    for elem in doc.element.body:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'p':
            # párrafo suelto
            from docx.text.paragraph import Paragraph as DocxParagraph
            parrs_con_tabla.append(('p', DocxParagraph(elem, doc)))
        elif tag == 'tbl':
            from docx.table import Table as DocxTable
            t = DocxTable(elem, doc)
            parrs_con_tabla.append(('tbl', t))

    # Modo simple: leer todos los párrafos en orden (incluyendo tablas)
    lineas = []
    for tipo, obj in parrs_con_tabla:
        if tipo == 'p':
            lineas.append(obj.text.strip())
        else:
            # tabla: leer celda por celda
            for row in obj.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = p.text.strip()
                        if txt:
                            lineas.append(txt)

    # Parsear preguntas
    i = 0
    primera_pregunta = None
    pregunta_actual = None
    alt_actual = []

    while i < len(lineas):
        linea = lineas[i]

        m_preg = RE_PREGUNTA.match(linea)
        m_alt  = RE_ALT.match(linea)

        if m_preg:
            # Guardar pregunta anterior
            if pregunta_actual is not None:
                pregunta_actual.alternativas = alt_actual[:]
                preguntas.append(pregunta_actual)

            if primera_pregunta is None:
                primera_pregunta = i

            num      = int(m_preg.group(1))
            enun     = m_preg.group(2).strip()
            pregunta_actual = Pregunta(num, enun, [])
            alt_actual = []

        elif m_alt and pregunta_actual is not None:
            letra = m_alt.group(1).upper()
            texto = m_alt.group(2).strip()
            alt_actual.append((letra, texto))

        elif pregunta_actual is None and linea:
            encabezado_parrs.append(linea)

        i += 1

    # Última pregunta
    if pregunta_actual is not None:
        pregunta_actual.alternativas = alt_actual[:]
        preguntas.append(pregunta_actual)

    return preguntas, encabezado_parrs


# ─────────────────────────────────────────────────────────────────────────────
# MEZCLADOR
# ─────────────────────────────────────────────────────────────────────────────

def mezclar_prueba(preguntas: list, semilla: int, mezclar_alts: bool = True) -> tuple:
    """
    Retorna (preguntas_mezcladas, mapa_equivalencias)
    mapa_equivalencias: lista de (num_original, letra_original, num_nueva, letra_nueva)
    """
    rng = random.Random(semilla)

    # Mezclar orden de preguntas
    indices = list(range(len(preguntas)))
    rng.shuffle(indices)

    nuevas = []
    mapa = []  # (num_A, letra_A, num_nueva, letra_nueva)

    for nuevo_num, idx_original in enumerate(indices, 1):
        p_orig = preguntas[idx_original]
        p_nueva = Pregunta(nuevo_num, p_orig.enunciado, [])

        if mezclar_alts and p_orig.alternativas:
            # Mezclar alternativas
            alts = list(p_orig.alternativas)
            rng.shuffle(alts)
            letras_nuevas = [chr(65 + i) for i in range(len(alts))]  # A, B, C...

            # Encontrar cuál letra corresponde a la correcta original
            letra_orig_correcta = p_orig.respuesta_correcta
            letra_nueva_correcta = None
            alts_renombradas = []
            for j, (letra_orig, texto) in enumerate(alts):
                nueva_letra = letras_nuevas[j]
                alts_renombradas.append((nueva_letra, texto))
                if letra_orig == letra_orig_correcta:
                    letra_nueva_correcta = nueva_letra

            p_nueva.alternativas = alts_renombradas
            p_nueva.respuesta_correcta = letra_nueva_correcta
        else:
            p_nueva.alternativas = list(p_orig.alternativas)
            p_nueva.respuesta_correcta = p_orig.respuesta_correcta

        nuevas.append(p_nueva)
        mapa.append({
            'num_A':       p_orig.numero,
            'letra_A':     p_orig.respuesta_correcta or '?',
            'num_nueva':   nuevo_num,
            'letra_nueva': p_nueva.respuesta_correcta or '?',
        })

    return nuevas, mapa


# ─────────────────────────────────────────────────────────────────────────────
# ESCRITOR: genera el .docx de una fila
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def escribir_docx_fila(preguntas: list, fila: str, encabezado_orig: list,
                        nombre_archivo: str, fila_original: str = "A"):
    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2)
        sec.right_margin  = Cm(2)

    # ── Encabezado ──────────────────────────────────────────────────────────
    enc = doc.add_paragraph()
    enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = enc.add_run("Departamento de Matemática")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AZUL

    # Tabla de identificación
    t_id = doc.add_table(rows=2, cols=3)
    t_id.style = 'Table Grid'
    hdrs = ["Nombre:", "Nivel:", f"Fila: {fila}"]
    from docx.shared import Pt as _Pt
    for i, h in enumerate(hdrs):
        cell = t_id.rows[0].cells[i]
        set_cell_bg(cell, "005A7F")
        p2 = cell.paragraphs[0]
        p2.clear()
        run = p2.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = _Pt(10)
        t_id.rows[1].cells[i].text = ""

    doc.add_paragraph()

    # Indicador de fila grande
    p_fila = doc.add_paragraph()
    p_fila.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_fila = p_fila.add_run(f"FILA {fila}")
    r_fila.bold = True
    r_fila.font.size = Pt(18)
    r_fila.font.color.rgb = AZUL

    doc.add_paragraph("Instrucciones: Lee cada pregunta con atención y marca la alternativa correcta.")
    doc.add_paragraph()

    # ── Preguntas ─────────────────────────────────────────────────────────────
    letras = ['a', 'b', 'c', 'd', 'e']

    for preg in preguntas:
        # Enunciado
        p_enun = doc.add_paragraph()
        r_num  = p_enun.add_run(f"{preg.numero}. ")
        r_num.bold = True
        r_num.font.size = Pt(11)
        r_txt = p_enun.add_run(preg.enunciado)
        r_txt.font.size = Pt(11)

        # Alternativas en tabla de 1 fila × 5 columnas (estilo horizontal)
        if preg.alternativas:
            n_alts = len(preg.alternativas)
            t_alts = doc.add_table(rows=1, cols=n_alts)
            t_alts.style = 'Table Grid'
            for j, (letra, texto) in enumerate(preg.alternativas):
                cell = t_alts.rows[0].cells[j]
                set_cell_bg(cell, "F5F5F5")
                par = cell.paragraphs[0]
                par.clear()
                r_l = par.add_run(f"{letra.lower()}) ")
                r_l.bold = True
                r_l.font.size = Pt(10)
                r_v = par.add_run(texto)
                r_v.font.size = Pt(10)

        doc.add_paragraph()  # espacio entre preguntas

    # ── Pie ──────────────────────────────────────────────────────────────────
    pie = doc.add_paragraph()
    r_pie = pie.add_run("proscar.cl — Generador de Filas")
    r_pie.font.size = Pt(8)
    r_pie.font.color.rgb = RGBColor(180, 180, 180)
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(nombre_archivo)
    print(f"  ✅ Fila {fila}: {nombre_archivo}")


# ─────────────────────────────────────────────────────────────────────────────
# TABLA DE EQUIVALENCIAS
# ─────────────────────────────────────────────────────────────────────────────

def escribir_tabla_equivalencias(mapas: dict, filas: list, n_preguntas: int,
                                  nombre_archivo: str):
    """
    mapas = { 'B': [...], 'C': [...] }
    Genera un .docx con la tabla de equivalencias para el docente.
    """
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2)
        sec.right_margin  = Cm(2)

    # Título
    tit = doc.add_paragraph()
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tit.add_run("TABLA DE EQUIVALENCIAS — USO DOCENTE")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AZUL

    doc.add_paragraph("Esta tabla permite corregir todas las filas usando la misma pauta de respuestas.")
    doc.add_paragraph()

    # Construir tabla
    # Columnas: Preg.A | Resp.A | Preg.B | Resp.B | Preg.C | Resp.C | ...
    n_cols = 2 + 2 * len(filas)
    tabla  = doc.add_table(rows=1 + n_preguntas, cols=n_cols)
    tabla.style = 'Table Grid'

    # Encabezados
    headers = ["Nº Fila A", "Resp. A"]
    for f in filas:
        headers += [f"Nº Fila {f}", f"Resp. {f}"]

    for j, h in enumerate(headers):
        cell = tabla.rows[0].cells[j]
        set_cell_bg(cell, "005A7F")
        par = cell.paragraphs[0]
        par.clear()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(h)
        run.bold = True
        run.font.color.rgb = BLANCO
        run.font.size = Pt(9)

    # Datos: iterar sobre las preguntas de fila A (1..n)
    # Para cada fila alternativa, buscar en el mapa cuál número de pregunta
    # corresponde a la pregunta i de fila A

    # Construir índice inverso para cada fila:
    # inv[fila][num_A] = (num_nueva, letra_nueva)
    inv = {}
    for f, mapa in mapas.items():
        inv[f] = {}
        for item in mapa:
            inv[f][item['num_A']] = (item['num_nueva'], item['letra_nueva'])

    # Respuestas fila A (asumimos que el usuario las indica o dejamos '?')
    # Aquí dejamos '?' para que el docente las complete
    for i in range(1, n_preguntas + 1):
        fila_tabla = tabla.rows[i]
        # Fila A
        fila_tabla.cells[0].text = str(i)
        fila_tabla.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        fila_tabla.cells[1].text = "___"
        fila_tabla.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for k, f in enumerate(filas):
            num_nueva, letra_nueva = inv[f].get(i, ('?', '?'))
            col_num  = 2 + 2*k
            col_resp = 3 + 2*k
            fila_tabla.cells[col_num].text  = str(num_nueva)
            fila_tabla.cells[col_num].paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.CENTER
            fila_tabla.cells[col_resp].text = letra_nueva if letra_nueva != '?' else "___"
            fila_tabla.cells[col_resp].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    nota = doc.add_paragraph()
    nota.add_run("Instrucciones: Complete la columna 'Resp. A' con las respuestas correctas de la fila A. "
                 "Las respuestas de las demás filas se calculan automáticamente según el orden de mezcla.")
    nota.runs[0].font.size = Pt(9)

    pie = doc.add_paragraph()
    pie.add_run("proscar.cl — Generador de Filas").font.size = Pt(8)
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(nombre_archivo)
    print(f"  ✅ Tabla de equivalencias: {nombre_archivo}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(
        description="Genera filas B, C, D de una prueba a partir de la fila A"
    )
    ap.add_argument("--input",    required=True, help="Prueba fila A (.docx)")
    ap.add_argument("--filas",    nargs="+", default=["B"], help="Filas a generar: B C D")
    ap.add_argument("--semilla",  type=int, default=2024, help="Semilla aleatoria base")
    ap.add_argument("--no-mezclar-alternativas", action="store_true",
                    help="Mezcla solo preguntas, no alternativas")
    ap.add_argument("--output-dir", default="", help="Carpeta de salida")
    args = ap.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"❌ No se encontró el archivo: {input_path}")
        return

    out_dir = Path(args.output_dir) if args.output_dir else OUT_DIR
    out_dir.mkdir(exist_ok=True)

    stem = input_path.stem

    print(f"\n📄 Leyendo: {input_path.name}")
    preguntas, encabezado = leer_preguntas_docx(str(input_path))

    if not preguntas:
        print("❌ No se encontraron preguntas en el documento.")
        print("   Asegúrate de que las preguntas estén numeradas: '1. Enunciado'")
        print("   Y las alternativas: 'a) opción' o 'A) opción'")
        return

    print(f"   → {len(preguntas)} preguntas detectadas")
    for p in preguntas:
        print(f"     P{p.numero}: {p.enunciado[:50]}... ({len(p.alternativas)} alts)")

    mezclar_alts = not args.no_mezclar_alternativas
    mapas = {}

    print(f"\n🔀 Generando filas: {', '.join(args.filas)}")

    for i, fila in enumerate(args.filas):
        # Semilla diferente por fila: base + índice
        semilla_fila = args.semilla + (i + 1) * 137
        pregs_mezcladas, mapa = mezclar_prueba(preguntas, semilla_fila, mezclar_alts)
        mapas[fila] = mapa

        nombre_out = str(out_dir / f"{stem}_Fila_{fila}.docx")
        escribir_docx_fila(pregs_mezcladas, fila, encabezado, nombre_out)

    # Tabla de equivalencias
    print(f"\n📊 Generando tabla de equivalencias...")
    nombre_equiv = str(out_dir / f"{stem}_Tabla_Equivalencias.docx")
    escribir_tabla_equivalencias(mapas, args.filas, len(preguntas), nombre_equiv)

    # Resumen JSON (útil para integración con otros scripts)
    resumen = {
        "archivo_original": str(input_path),
        "n_preguntas": len(preguntas),
        "filas_generadas": args.filas,
        "semilla_base": args.semilla,
        "mezcladas_alternativas": mezclar_alts,
        "mapas": {f: mapas[f] for f in args.filas}
    }
    json_path = out_dir / f"{stem}_equivalencias.json"
    json_path.write_text(
        json.dumps(resumen, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )
    print(f"  ✅ JSON de equivalencias: {json_path}")

    print(f"\n✅ Listo. Archivos generados en: {out_dir}/")
    print(f"   • {len(args.filas)} fila(s) nueva(s)")
    print(f"   • 1 tabla de equivalencias para el docente")

if __name__ == "__main__":
    main()
